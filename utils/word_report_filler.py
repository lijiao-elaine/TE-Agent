from docx import Document
from docx.table import Table
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from typing import Dict, List, Any
import os
from datetime import datetime
import pdb
# 兼容不同版本的python-docx库
try:
    # 旧版本导入方式
    from docx.oxml.xmlchemy import OxmlElement
except ImportError:
    # 新版本导入方式
    from docx.oxml import OxmlElement

class WordReportFiller:
    """处理测试结果向Word文档的填充（基于测试用例表格）"""

    @staticmethod
    def find_case_table(doc: Document, case_id: str) -> Table:
        """
        根据用例标识查找对应的测试表格（适配merged_document.docx结构）
        :param doc: 打开的Word文档对象
        :param case_id: 测试用例标识（如XXX_TEST_001）
        :return: 匹配的表格对象
        """
        #pdb.set_trace()
        for table_idx, table in enumerate(doc.tables):
            #print(f"\n===== 查找第 {table_idx + 1} 个用例表格是否是case_id为{case_id}的用例 =====")

            # 遍历表格的每一行
            for row_idx, row in enumerate(table.rows):
                # 遍历行中的每一列（单元格）
                for col_idx, cell in enumerate(row.cells):
                    # 获取单元格文本（处理空单元格的情况）
                    cell_text = "\n".join([para.text for para in cell.paragraphs]).strip()
                    #if cell_text:  # 只处理非空单元格
                    #    print(f"表格{table_idx + 1} 行{row_idx + 1} 列{col_idx + 1}: {cell_text}")
                    if case_id in cell_text:
                        #print(f"找到用例：{case_id}，in 表格{table_idx + 1} 行{row_idx + 1} 列{col_idx + 1}: {cell_text}")
                        return table
        raise ValueError(f"在文档中未找到用例标识为 {case_id} 的表格")

    @staticmethod
    def insert_images_after_placeholder(cell, image_paths: List[str], placeholder: str = "其它____", max_width: float = 2.0):
        """在占位符后追加内容并插入图片"""
        # 确保单元格至少有一个段落
        if not cell.paragraphs:
            cell.add_paragraph()
        
        # 1. 查找占位符所在段落
        placeholder_para = None
        for para in cell.paragraphs:
            if placeholder in para.text:
                placeholder_para = para
                break
        
        # 如未找到占位符，添加一个
        if not placeholder_para:
            placeholder_para = cell.add_paragraph(placeholder)
        
        # 2. 获取占位符段落的XML元素
        placeholder_elem = placeholder_para._element
        
        # 3. 逐个插入图片（每个图片作为新段落插入到占位符之后）
        for img_idx, img_path in enumerate(image_paths, 1):
            try:
                # 创建新段落XML元素
                new_para_elem = OxmlElement('w:p')
                
                # 在占位符元素后面插入新段落，确保图片顺序正确
                placeholder_elem.addnext(new_para_elem)
                
                # 创建段落对象
                new_para = Paragraph(new_para_elem, cell)
                
                # 处理图片不存在的情况
                if not os.path.exists(img_path):
                    new_para.text = f"[图片不存在: {os.path.basename(img_path)}]"
                    if new_para.runs:
                        new_para.runs[0].font.color.rgb = (0xFF, 0x00, 0x00)
                    continue
                
                # 添加图片标题和图片
                run = new_para.add_run(f"截图 {img_idx}: ")
                run.bold = True
                run = new_para.add_run()
                run.add_picture(img_path, width=Inches(max_width))
                
                # 更新占位符元素为刚插入的段落，确保下一张图片插在它后面
                placeholder_elem = new_para_elem
                
            except Exception as e:
                # 异常处理
                error_para = cell.add_paragraph()
                error_para.text = f"[插入图片失败: {os.path.basename(img_path)}，错误: {str(e)}]"
                error_para.runs[0].font.color.rgb = (0xFF, 0x00, 0x00)
                # 更新占位符元素，避免影响后续图片
                placeholder_elem = error_para._element

    @staticmethod
    def fill_case_results_old(word_file: str, case_result: Dict[str, Any]) -> bool:
        """将测试结果填充到Word文档的对应表格中"""
        doc = Document(word_file)
        
        try:
            table = WordReportFiller.find_case_table(doc, case_result["case_id"])
            step_results = case_result["execution_steps"]
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = "\n".join([para.text for para in cell.paragraphs]).strip()                 
                    # 处理步骤结果（通过/不通过）
                    if "□通过" in cell_text and "□不通过" in cell_text and "其它____" in cell_text:
                        step_idx = int(row.cells[0].text) - 1
                        
                        # 更新通过/不通过状态
                        if step_results[step_idx]["step_result"] == "通过":
                            for para in cell.paragraphs:
                                para.text = para.text.replace("□通过", "☑通过")
                        else:
                            for para in cell.paragraphs:
                                para.text = para.text.replace("□不通过", "☑不通过")
                        #print(f"步骤{row.cells[0].text}的执行结果，回填入用例表格第{row_idx + 1}行第{col_idx + 1}列")

                        # 回填每个步骤的截图, 在"其它____"后追加
                        screenshot_paths = step_results[step_idx].get("screenshot_path", [])
                        if screenshot_paths and isinstance(screenshot_paths, list):
                            WordReportFiller.insert_images_after_placeholder(
                                cell, 
                                screenshot_paths,
                                placeholder="其它____",
                                max_width=1.0  # 图片最大宽度
                            )
                    
                    # 记录总体结果位置
                    if cell_text == "测试用例执行结果":
                        overall_row_idx = row_idx
                        overall_col_idx = col_idx 
                    # 记录测试时间、测试人员、操作人员位置
                    elif cell_text == "测试时间":
                        test_time_row_idx = row_idx
                        test_time_col_idx = col_idx 
                    elif cell_text == "测试人员":
                        test_person_row_idx = row_idx
                        test_person_col_idx = col_idx 
                    elif cell_text == "操作人员":
                        operate_person_row_idx = row_idx
                        operate_person_col_idx = col_idx
                    else:
                        continue
            
            # 回填总体结果
            if case_result["overall_result"] == "通过":
                table.rows[overall_row_idx].cells[overall_col_idx + 1].text = "通过"
            else:
                table.rows[overall_row_idx].cells[overall_col_idx + 1].text = "不通过"
            # 回填测试时间、测试人员、操作人员
            test_time = datetime.now()
            table.rows[test_time_row_idx].cells[test_time_col_idx + 1].text = test_time.strftime("%Y-%m-%d %H:%M:%S")
            table.rows[test_person_row_idx].cells[test_person_col_idx + 1].text = "auto run"
            table.rows[operate_person_row_idx].cells[operate_person_col_idx + 1].text = "auto run"
            
            doc.save(word_file)
            return True
        except Exception as e:
            raise RuntimeError(f"回填Word报告失败: {str(e)}")
            return False

    @staticmethod
    def fill_case_results(word_file: str, step_num: int, case_result: Dict[str, Any]) -> bool:
        """将测试结果填充到Word文档的对应表格中"""
        doc = Document(word_file)
        
        try:
            table = WordReportFiller.find_case_table(doc, case_result["case_id"])
            step_results = case_result["execution_steps"]
            result_len = len(step_results)
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = "\n".join([para.text for para in cell.paragraphs]).strip()                 
                    # 处理步骤结果（通过/不通过）
                    if "□通过" in cell_text and "□不通过" in cell_text and "其它____" in cell_text:
                        step_idx = int(row.cells[0].text) - 1
                        if result_len > 0: # run_test_step 中，执行case_result["steps"].append前发生了Exception，结果数比执行步数少，只回填有结果的步骤
                            if step_num == -1: # 预处理失败
                                pass
                            elif step_idx + 1 <= step_num: # 预处理成功，但仅执行了 step_num 个测试步骤，也只回填这些步骤的结果
                                if step_results[step_idx]["step_result"] == "通过":
                                    for para in cell.paragraphs:
                                        para.text = para.text.replace("□通过", "☑通过")
                                else:
                                    for para in cell.paragraphs:
                                        para.text = para.text.replace("□不通过", "☑不通过")
                                #print(f"步骤{row.cells[0].text}的执行结果，回填入用例表格第{row_idx + 1}行第{col_idx + 1}列")

                                # 回填每个步骤的截图, 在"其它____"后追加
                                screenshot_paths = step_results[step_idx].get("screenshot_path", [])
                                if screenshot_paths and isinstance(screenshot_paths, list):
                                    WordReportFiller.insert_images_after_placeholder(
                                        cell, 
                                        screenshot_paths,
                                        placeholder="其它____",
                                        max_width=1.0  # 图片最大宽度
                                    )
                            else:
                                pass
                            result_len -= 1
                    # 记录总体结果位置
                    if cell_text == "测试用例执行结果":
                        overall_row_idx = row_idx
                        overall_col_idx = col_idx 
                    # 记录测试时间、测试人员、操作人员位置
                    elif cell_text == "测试时间":
                        test_time_row_idx = row_idx
                        test_time_col_idx = col_idx 
                    elif cell_text == "测试人员":
                        test_person_row_idx = row_idx
                        test_person_col_idx = col_idx 
                    elif cell_text == "操作人员":
                        operate_person_row_idx = row_idx
                        operate_person_col_idx = col_idx
                    else:
                        continue
            
            # 回填总体结果、测试时间、测试人员、操作人员
            table.rows[overall_row_idx].cells[overall_col_idx + 1].text = case_result["overall_result"]
            test_time = datetime.now()
            table.rows[test_time_row_idx].cells[test_time_col_idx + 1].text = test_time.strftime("%Y-%m-%d %H:%M:%S")
            table.rows[test_person_row_idx].cells[test_person_col_idx + 1].text = "auto run"
            table.rows[operate_person_row_idx].cells[operate_person_col_idx + 1].text = "auto run"
            
            doc.save(word_file)
            return True
        except Exception as e:
            raise RuntimeError(f"回填Word报告失败: {str(e)}")
            return False

